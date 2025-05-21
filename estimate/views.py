from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib.auth import login, authenticate
from django.contrib import messages
from django.http import HttpResponse
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import _Cell
import os
from datetime import datetime
from .models import Estimate, PaverBlockType
from .forms import CustomLoginForm, EstimateForm, PaverBlockTypeForm
import tempfile
import logging
import traceback
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import docx2txt

logger = logging.getLogger(__name__)

def login_view(request):
    if request.method == 'POST':
        form = CustomLoginForm(data=request.POST)
        if form.is_valid():
            username = form.cleaned_data.get('username')
            password = form.cleaned_data.get('password')
            user = authenticate(username=username, password=password)
            if user is not None:
                login(request, user)
                return redirect('dashboard')
    else:
        form = CustomLoginForm()
    return render(request, 'estimate/login.html', {'form': form})

@login_required
def dashboard(request):
    estimates = Estimate.objects.filter(created_by=request.user).order_by('-created_at')
    return render(request, 'estimate/dashboard.html', {'estimates': estimates})

@login_required
def create_estimate(request):
    if request.method == 'POST':
        form = EstimateForm(request.POST)
        if form.is_valid():
            estimate = form.save(commit=False)
            estimate.created_by = request.user
            estimate.save()
            messages.success(request, 'Estimate created successfully!')
            return redirect('dashboard')
    else:
        form = EstimateForm()
    return render(request, 'estimate/create_estimate.html', {'form': form})

@login_required
def manage_paver_blocks(request):
    if request.method == 'POST':
        form = PaverBlockTypeForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'Paver block type added successfully!')
            return redirect('manage_paver_blocks')
    else:
        form = PaverBlockTypeForm()
    
    paver_blocks = PaverBlockType.objects.all()
    return render(request, 'estimate/manage_paver_blocks.html', {
        'form': form,
        'paver_blocks': paver_blocks
    })

@login_required
def delete_paver_block(request, paver_block_id):
    paver_block = get_object_or_404(PaverBlockType, id=paver_block_id)
    if request.method == 'POST':
        paver_block.delete()
        messages.success(request, 'Paver block type deleted successfully!')
        return redirect('manage_paver_blocks')
    return render(request, 'estimate/confirm_delete_paver_block.html', {'paver_block': paver_block})

def replace_text_in_docx(docx_path, replacements):
    """Replace text in a DOCX file with given replacements."""
    doc = Document(docx_path)
    
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, str(value))
    
    return doc

def create_pdf_from_docx(docx_path, replacements):
    """Create a PDF from a DOCX file with replacements."""
    # Extract text from DOCX
    text = docx2txt.process(docx_path)
    
    # Apply replacements
    for key, value in replacements.items():
        text = text.replace(key, str(value))
    
    # Create PDF
    buffer = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
    pdf_path = buffer.name
    buffer.close()
    
    doc = SimpleDocTemplate(
        pdf_path,
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72
    )
    
    # Create styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=30
    )
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=12,
        spaceAfter=12
    )
    
    # Create content
    content = []
    
    # Split text into paragraphs and create PDF content
    paragraphs = text.split('\n\n')
    for para in paragraphs:
        if para.strip():
            # Check if it's a title (you might want to adjust this logic)
            if para.strip().isupper() or para.strip().startswith('KCP'):
                content.append(Paragraph(para.strip(), title_style))
            else:
                content.append(Paragraph(para.strip(), normal_style))
            content.append(Spacer(1, 12))
    
    # Build PDF
    doc.build(content)
    return pdf_path

@login_required
def generate_pdf(request, estimate_id):
    try:
        estimate = get_object_or_404(Estimate, id=estimate_id, created_by=request.user)
        template_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'KCP_LETTERPAD.docx')
        
        if not os.path.exists(template_path):
            logger.error(f"Template file not found at: {template_path}")
            messages.error(request, 'Template file not found. Please contact support.')
            return redirect('dashboard')
            
        # Create temporary files
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as docx_file:
            docx_filename = docx_file.name
            
            # Copy template to temp file
            with open(template_path, 'rb') as src:
                with open(docx_filename, 'wb') as dst:
                    dst.write(src.read())
            
            logger.info(f"Copied template to: {docx_filename}")
            
            # Replace placeholders
            replacements = {
                '<partyname>': estimate.party_name,
                '<date>': str(estimate.date),
                '<paverblocktype>': str(estimate.paver_block_type),
                '<rate1>': str(estimate.price),
                '<rate2>': str(estimate.gst_amount),
                '<rate3>': str(estimate.transportation_charge),
                '<rate4>': str(estimate.loading_unloading_cost),
                '<rate5>': str(estimate.loading_unloading_cost),
                '<rate>': str(estimate.total_amount),
                '<year>': str(datetime.now().year),
                '<NOTE>': estimate.notes or '',
            }
            
            # Replace placeholders in the document
            doc = replace_text_in_docx(docx_filename, replacements)
            doc.save(docx_filename)
            logger.info("Replaced placeholders in document")

        try:
            # Convert to PDF using reportlab
            logger.info("Attempting PDF conversion with reportlab")
            pdf_filename = create_pdf_from_docx(docx_filename, replacements)
            
            if os.path.exists(pdf_filename):
                logger.info(f"PDF file created successfully at: {pdf_filename}")
                with open(pdf_filename, 'rb') as f:
                    response = HttpResponse(f.read(), content_type='application/pdf')
                    response['Content-Disposition'] = f'attachment; filename="KCP-ESTIMATE-{estimate.party_name}.pdf"'
                os.remove(pdf_filename)
                os.remove(docx_filename)
                return response
            else:
                logger.error(f"PDF file not created at expected location: {pdf_filename}")
                raise FileNotFoundError("PDF file was not created")
                
        except Exception as e:
            logger.error(f"PDF conversion failed: {str(e)}")
            logger.error(f"Traceback: {traceback.format_exc()}")
            # If PDF conversion fails, return the DOCX file
            with open(docx_filename, 'rb') as f:
                response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = f'attachment; filename="KCP-ESTIMATE-{estimate.party_name}.docx"'
            os.remove(docx_filename)
            messages.warning(request, 'PDF conversion failed. Downloading DOCX file instead.')
            return response

    except Exception as e:
        logger.error(f"Error in generate_pdf: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        messages.error(request, f'Error generating document: {str(e)}')
        return redirect('dashboard')

@login_required
def delete_estimate(request, estimate_id):
    estimate = get_object_or_404(Estimate, id=estimate_id, created_by=request.user)
    if request.method == 'POST':
        estimate.delete()
        messages.success(request, 'Estimate deleted successfully!')
        return redirect('dashboard')
    # Optional: Render a confirmation page for GET requests
    return render(request, 'estimate/confirm_delete.html', {'estimate': estimate})
